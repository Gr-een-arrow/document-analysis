import os
import tempfile
from django.test import TestCase
from services.parser import DocumentParser

class TestDocumentParser(TestCase):
    def test_parse_txt(self):
        content = "Hello\nWorld"
        with tempfile.NamedTemporaryFile(mode='w+', suffix='.txt', delete=False) as f:
            f.write(content)
            f.flush()
            result = DocumentParser.parse_txt(f.name)
        self.assertEqual(result['text'], content)
        self.assertEqual(result['num_lines'], 2)
        os.remove(f.name)

    def test_parse_docx(self):
        import docx
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        try:
            doc = docx.Document()
            doc.add_paragraph("Hello")
            doc.add_paragraph("World")
            doc.save(tmp.name)
            result = DocumentParser.parse_docx(tmp.name)
            self.assertIn("Hello", result['text'])
            self.assertIn("World", result['text'])
            self.assertEqual(result['num_paragraphs'], 2)
        finally:
            tmp.close()
            os.remove(tmp.name)

    def test_parse_pdf(self):
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen.canvas import Canvas
        tmp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        try:
            c = Canvas(tmp.name, pagesize=letter)
            c.drawString(100, 750, "Hello World")
            c.save()
            result = DocumentParser.parse_pdf(tmp.name)
            self.assertIn("Hello World", result['text'])
            self.assertEqual(result['num_pages'], 1)
        finally:
            tmp.close()
            os.remove(tmp.name)

    def test_parse_excel(self):
        import openpyxl
        tmp = tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False)
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws['A1'] = 'Hello'
            ws['B1'] = 'World'
            wb.save(tmp.name)
            wb.close()
            result = DocumentParser.parse_excel(tmp.name)
            self.assertIn("Hello", result['text'])
            self.assertIn("World", result['text'])
            self.assertEqual(result['num_sheets'], 1)
        finally:
            tmp.close()
            os.remove(tmp.name)
