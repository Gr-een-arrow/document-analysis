import os
from typing import List, Dict, Any

from PyPDF2 import PdfReader
import docx
import openpyxl

class DocumentParser:
    """
    Utility class for parsing PDF, DOCX, TXT, and Excel files.
    Returns extracted text and basic metadata.
    """

    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return {"text": text, "num_pages": len(reader.pages)}

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return {"text": text, "num_paragraphs": len(doc.paragraphs)}

    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {"text": text, "num_lines": text.count('\n') + 1}

    @staticmethod
    def parse_excel(file_path: str) -> Dict[str, Any]:
        wb = openpyxl.load_workbook(file_path, read_only=True)
        text = []
        for sheet in wb:
            for row in sheet.iter_rows(values_only=True):
                text.append("\t".join([str(cell) if cell is not None else '' for cell in row]))
        wb.close()
        return {"text": "\n".join(text), "num_sheets": len(wb.sheetnames)}

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return DocumentParser.parse_pdf(file_path)
        elif ext == '.docx':
            return DocumentParser.parse_docx(file_path)
        elif ext == '.txt':
            return DocumentParser.parse_txt(file_path)
        elif ext in ['.xls', '.xlsx']:
            return DocumentParser.parse_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
